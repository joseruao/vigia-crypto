from __future__ import annotations

import hmac
import io
import logging
import os
from pathlib import Path

from fastapi import APIRouter, File, Form, Header, HTTPException, UploadFile

from Api.services.pme_procurement import (
    ProcurementAnalysisResponse,
    RawCatalog,
    compare_catalogs,
    parse_commercial_values,
    parse_purchase_needs,
)

log = logging.getLogger("vigia.pme")
router = APIRouter(prefix="/api/pme", tags=["pme"])

SUPPORTED_TEXT_SUFFIXES = {".csv", ".tsv", ".txt"}
SUPPORTED_DOCUMENT_SUFFIXES = {".pdf", ".docx"}
MAX_CATALOG_BYTES = 4 * 1024 * 1024


def _check_access_code(provided: str | None) -> None:
    if os.getenv("PME_REQUIRE_ACCESS_CODE", "true").strip().lower() in {"0", "false", "no"}:
        return
    expected = os.getenv("PME_ACCESS_CODE", "") or os.getenv("DEVILS_ADVOCATE_ACCESS_CODE", "")
    if not expected:
        is_production = bool(os.getenv("RENDER") or os.getenv("RENDER_SERVICE_ID")) or "joseruao.com" in os.getenv("FRONTEND_URL", "")
        if not is_production:
            return
        raise HTTPException(
            status_code=503,
            detail="PME is not configured for access. Set PME_ACCESS_CODE.",
        )
    if not provided or not hmac.compare_digest(provided.strip(), expected):
        raise HTTPException(status_code=401, detail="Invalid or missing access code.")


async def _read_catalog(file: UploadFile) -> RawCatalog:
    filename = file.filename or "catalogo.csv"
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_TEXT_SUFFIXES and suffix not in SUPPORTED_DOCUMENT_SUFFIXES:
        raise ValueError(
            f"{filename}: use CSV, TSV, TXT, PDF textual ou DOCX. Imagens/scans precisam de OCR na próxima fase."
        )
    content = await file.read(MAX_CATALOG_BYTES + 1)
    if len(content) > MAX_CATALOG_BYTES:
        raise ValueError(f"{filename}: ficheiro demasiado grande. Limite: 4 MB.")
    if suffix == ".pdf":
        return RawCatalog(filename=filename, text=_extract_pdf_text(content))
    if suffix == ".docx":
        return RawCatalog(filename=filename, text=_extract_docx_text(content))
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return RawCatalog(filename=filename, text=content.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise ValueError(f"{filename}: não consegui ler o texto do ficheiro.")


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF extraction requires pypdf.") from exc

    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages[:80])


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires python-docx.") from exc

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@router.post("/procurement/analyze", response_model=ProcurementAnalysisResponse)
async def analyze_procurement(
    files: list[UploadFile] = File(...),
    needs_text: str = Form(default=""),
    commercial_values_text: str = Form(default=""),
    x_access_code: str | None = Header(default=None),
):
    _check_access_code(x_access_code)
    if not files:
        raise HTTPException(status_code=400, detail="Carregue pelo menos um catálogo.")
    if len(files) > 12:
        raise HTTPException(status_code=400, detail="Máximo de 12 catálogos por análise.")
    try:
        catalogs = [await _read_catalog(file) for file in files]
        return compare_catalogs(
            catalogs,
            parse_purchase_needs(needs_text),
            parse_commercial_values(commercial_values_text),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        log.error("pme procurement analysis failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail="A análise de compras falhou.") from exc
